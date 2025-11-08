"""
Solar PV Reliability Test Lab SOP Generator - Complete Application
Integrates all existing modules with comprehensive SOP generation capabilities
"""

import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime, timedelta, date
import uuid
import io
import base64
from typing import List, Dict, Any
import json
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from PIL import Image
import networkx as nx

# ============================================================================
# PAGE CONFIGURATION
# ============================================================================

st.set_page_config(
    page_title="Solar PV SOP Generator",
    page_icon="☀️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================================
# CUSTOM CSS FOR PROFESSIONAL UI
# ============================================================================

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #FF6B35;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #FF6B35 0%, #F7931E 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .sub-header {
        font-size: 1.5rem;
        color: #004E89;
        font-weight: 600;
        padding: 0.5rem 0;
        border-bottom: 2px solid #FF6B35;
    }
    .info-box {
        background-color: #E8F4F8;
        border-left: 4px solid #004E89;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #D4EDDA;
        border-left: 4px solid #28A745;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #FFF3CD;
        border-left: 4px solid #FFC107;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 0.5rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        border-left: 4px solid #FF6B35;
    }
    .stButton>button {
        background-color: #FF6B35;
        color: white;
        font-weight: bold;
        border-radius: 0.5rem;
        padding: 0.5rem 2rem;
        border: none;
    }
    .stButton>button:hover {
        background-color: #E55A2B;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

def init_session_state():
    """Initialize all session state variables"""
    if 'initialized' not in st.session_state:
        st.session_state.initialized = True
        st.session_state.current_user = 'Test Engineer'
        st.session_state.user_role = 'Engineer'

        # Core data structures from existing modules
        st.session_state.projects = []
        st.session_state.tasks = []
        st.session_state.samples = []
        st.session_state.equipment = []
        st.session_state.manpower = []
        st.session_state.test_methods = []
        st.session_state.test_results = []
        st.session_state.risks = []
        st.session_state.issues = []
        st.session_state.approvals = []
        st.session_state.documents = []
        st.session_state.notifications = []
        st.session_state.signatures = []
        st.session_state.audit_trail = []
        st.session_state.holidays = []

        # SOP Generator specific
        st.session_state.sop_documents = []
        st.session_state.sop_templates = []
        st.session_state.standards_library = []
        st.session_state.uploaded_images = {}
        st.session_state.uploaded_diagrams = {}
        st.session_state.uploaded_tables = {}

        # Initialize with sample data
        init_sample_data()
        init_sop_templates()
        init_standards_library()

def init_sample_data():
    """Initialize sample data for all modules"""
    # Projects
    st.session_state.projects = [
        {
            'id': 'PRJ001',
            'name': 'IEC 61215 Module Qualification',
            'client': 'SolarTech Industries',
            'start_date': date(2024, 1, 1),
            'end_date': date(2024, 12, 31),
            'status': 'Active',
            'priority': 'High',
            'budget': 500000,
            'spent': 125000,
            'manager': 'Dr. Rajesh Kumar',
            'description': 'Full IEC 61215 qualification testing for crystalline silicon PV modules'
        },
        {
            'id': 'PRJ002',
            'name': 'IEC 61730 Safety Testing',
            'client': 'Green Energy Corp',
            'start_date': date(2024, 2, 1),
            'end_date': date(2024, 11, 30),
            'status': 'Active',
            'priority': 'High',
            'budget': 350000,
            'spent': 87500,
            'manager': 'Dr. Priya Sharma',
            'description': 'Safety qualification testing for photovoltaic modules'
        }
    ]

    # Equipment
    st.session_state.equipment = [
        {
            'id': 'EQP001',
            'name': 'Solar Simulator AAA Class',
            'type': 'Testing Equipment',
            'manufacturer': 'Spectra Nova',
            'model': 'SN-4000',
            'status': 'Available',
            'location': 'Lab Room 101',
            'last_calibration': date(2024, 10, 1),
            'next_calibration': date(2025, 4, 1),
            'performance_metric': 98.5,
            'utilization': 65,
            'accuracy': '±2%',
            'range': '0-1400 W/m²'
        },
        {
            'id': 'EQP002',
            'name': 'Thermal Cycling Chamber',
            'type': 'Environmental Chamber',
            'manufacturer': 'EnviroTest',
            'model': 'TC-2000',
            'status': 'In Use',
            'location': 'Lab Room 102',
            'last_calibration': date(2024, 9, 15),
            'next_calibration': date(2025, 3, 15),
            'performance_metric': 96.2,
            'utilization': 80,
            'temp_range': '-40°C to +85°C',
            'humidity_range': '10% to 85% RH'
        },
        {
            'id': 'EQP003',
            'name': 'UV Weathering Chamber',
            'type': 'Environmental Chamber',
            'manufacturer': 'WeatherPro',
            'model': 'UV-1000',
            'status': 'Available',
            'location': 'Lab Room 103',
            'last_calibration': date(2024, 11, 1),
            'next_calibration': date(2025, 5, 1),
            'performance_metric': 97.8,
            'utilization': 55,
            'uv_intensity': '15-60 W/m² at 340nm'
        }
    ]

    # Manpower
    st.session_state.manpower = [
        {
            'id': 'EMP001',
            'name': 'Arun Patel',
            'role': 'Senior Test Engineer',
            'skills': ['IEC 61215', 'IEC 61730', 'Thermal Testing'],
            'certifications': ['ISO 17025 Lead Assessor', 'IEC 61215 Certified'],
            'availability': 'Available',
            'performance_score': 92,
            'hours_logged': 1250,
            'email': 'arun.patel@lab.com',
            'phone': '+91-9876543210'
        },
        {
            'id': 'EMP002',
            'name': 'Sneha Reddy',
            'role': 'Test Engineer',
            'skills': ['Electrical Testing', 'Data Analysis', 'Report Writing'],
            'certifications': ['IEC 61730 Certified', 'NABL Trained'],
            'availability': 'Busy',
            'performance_score': 95,
            'hours_logged': 1480,
            'email': 'sneha.reddy@lab.com',
            'phone': '+91-9876543211'
        },
        {
            'id': 'EMP003',
            'name': 'Vikram Singh',
            'role': 'Lab Technician',
            'skills': ['Sample Preparation', 'Equipment Operation', 'Calibration'],
            'certifications': ['NABL Trained'],
            'availability': 'Available',
            'performance_score': 88,
            'hours_logged': 980,
            'email': 'vikram.singh@lab.com',
            'phone': '+91-9876543212'
        }
    ]

    # Test Methods
    st.session_state.test_methods = [
        {
            'id': 'TM001',
            'name': 'Thermal Cycling Test',
            'standard': 'IEC 61215-2:2016',
            'method_number': 'MQT 12',
            'category': 'Environmental',
            'equipment_required': ['Thermal Cycling Chamber', 'Data Logger'],
            'duration_hours': 200,
            'parameters': ['Temperature', 'Cycles', 'Humidity'],
            'acceptance_criteria': 'No visual defects, Pmax degradation < 5%'
        },
        {
            'id': 'TM002',
            'name': 'Humidity Freeze Test',
            'standard': 'IEC 61215-2:2016',
            'method_number': 'MQT 13',
            'category': 'Environmental',
            'equipment_required': ['Climate Chamber', 'IV Curve Tracer'],
            'duration_hours': 240,
            'parameters': ['Temperature', 'Humidity', 'Cycles'],
            'acceptance_criteria': 'No visual defects, Pmax degradation < 5%'
        },
        {
            'id': 'TM003',
            'name': 'UV Preconditioning Test',
            'standard': 'IEC 61215-2:2016',
            'method_number': 'MQT 09',
            'category': 'Environmental',
            'equipment_required': ['UV Chamber', 'Spectroradiometer'],
            'duration_hours': 168,
            'parameters': ['UV Dose', 'Temperature', 'Irradiance'],
            'acceptance_criteria': 'Pmax degradation < 5%'
        }
    ]

    # Samples
    st.session_state.samples = [
        {
            'id': 'SMP001',
            'name': 'Mono-Si Module 400W',
            'type': 'PV Module',
            'batch': 'BATCH-2024-001',
            'received_date': date(2024, 1, 15),
            'condition': 'New',
            'location': 'Storage Room A',
            'status': 'Testing',
            'barcode': 'MOD400W-001',
            'manufacturer': 'SolarTech',
            'power_rating': '400W',
            'dimensions': '1956 x 992 x 40 mm'
        },
        {
            'id': 'SMP002',
            'name': 'Poly-Si Module 350W',
            'type': 'PV Module',
            'batch': 'BATCH-2024-002',
            'received_date': date(2024, 1, 20),
            'condition': 'New',
            'location': 'Lab Room 101',
            'status': 'Conditioning',
            'barcode': 'MOD350W-002',
            'manufacturer': 'Green Energy',
            'power_rating': '350W',
            'dimensions': '1686 x 992 x 35 mm'
        }
    ]

    # Risks
    st.session_state.risks = [
        {
            'id': 'RSK001',
            'title': 'Equipment Calibration Delay',
            'category': 'Technical',
            'probability': 'Medium',
            'impact': 'High',
            'status': 'Active',
            'mitigation': 'Schedule calibration 2 months in advance, maintain backup equipment',
            'owner': 'Dr. Rajesh Kumar',
            'identified_date': date(2024, 1, 10)
        }
    ]

    # Holidays
    st.session_state.holidays = [
        {'date': date(2024, 1, 26), 'name': 'Republic Day', 'type': 'National'},
        {'date': date(2024, 8, 15), 'name': 'Independence Day', 'type': 'National'},
        {'date': date(2024, 10, 2), 'name': 'Gandhi Jayanti', 'type': 'National'},
    ]

def init_sop_templates():
    """Initialize SOP templates for different test types"""
    st.session_state.sop_templates = [
        {
            'id': 'TPL001',
            'name': 'Thermal Cycling Test SOP',
            'standard': 'IEC 61215-2:2016 MQT 12',
            'test_type': 'Environmental',
            'sections': ['Purpose', 'Scope', 'Definitions', 'Responsibilities', 'Equipment', 'Procedure', 'Analysis', 'Criteria']
        },
        {
            'id': 'TPL002',
            'name': 'Humidity Freeze Test SOP',
            'standard': 'IEC 61215-2:2016 MQT 13',
            'test_type': 'Environmental',
            'sections': ['Purpose', 'Scope', 'Definitions', 'Responsibilities', 'Equipment', 'Procedure', 'Analysis', 'Criteria']
        },
        {
            'id': 'TPL003',
            'name': 'UV Preconditioning Test SOP',
            'standard': 'IEC 61215-2:2016 MQT 09',
            'test_type': 'Environmental',
            'sections': ['Purpose', 'Scope', 'Definitions', 'Responsibilities', 'Equipment', 'Procedure', 'Analysis', 'Criteria']
        },
        {
            'id': 'TPL004',
            'name': 'Mechanical Load Test SOP',
            'standard': 'IEC 61215-2:2016 MQT 16',
            'test_type': 'Mechanical',
            'sections': ['Purpose', 'Scope', 'Definitions', 'Responsibilities', 'Equipment', 'Procedure', 'Analysis', 'Criteria']
        }
    ]

def init_standards_library():
    """Initialize standards library with common test standards"""
    st.session_state.standards_library = [
        {
            'id': 'STD001',
            'code': 'IEC 61215-1:2016',
            'title': 'Terrestrial photovoltaic (PV) modules - Design qualification and type approval - Part 1: Test requirements',
            'category': 'PV Module Testing',
            'authority': 'IEC',
            'year': 2016,
            'citation': 'IEC 61215-1:2016, Terrestrial photovoltaic (PV) modules - Design qualification and type approval - Part 1: Test requirements'
        },
        {
            'id': 'STD002',
            'code': 'IEC 61215-2:2016',
            'title': 'Terrestrial photovoltaic (PV) modules - Design qualification and type approval - Part 2: Test procedures',
            'category': 'PV Module Testing',
            'authority': 'IEC',
            'year': 2016,
            'citation': 'IEC 61215-2:2016, Terrestrial photovoltaic (PV) modules - Design qualification and type approval - Part 2: Test procedures'
        },
        {
            'id': 'STD003',
            'code': 'IEC 61730-1:2016',
            'title': 'Photovoltaic (PV) module safety qualification - Part 1: Requirements for construction',
            'category': 'Safety',
            'authority': 'IEC',
            'year': 2016,
            'citation': 'IEC 61730-1:2016, Photovoltaic (PV) module safety qualification - Part 1: Requirements for construction'
        },
        {
            'id': 'STD004',
            'code': 'IEC 61730-2:2016',
            'title': 'Photovoltaic (PV) module safety qualification - Part 2: Requirements for testing',
            'category': 'Safety',
            'authority': 'IEC',
            'year': 2016,
            'citation': 'IEC 61730-2:2016, Photovoltaic (PV) module safety qualification - Part 2: Requirements for testing'
        },
        {
            'id': 'STD005',
            'code': 'IS 14286:2005',
            'title': 'Crystalline Silicon Terrestrial Photovoltaic (PV) Modules - Design Qualification and Type Approval',
            'category': 'Indian Standard',
            'authority': 'BIS',
            'year': 2005,
            'citation': 'IS 14286:2005, Crystalline Silicon Terrestrial Photovoltaic (PV) Modules - Design Qualification and Type Approval'
        },
        {
            'id': 'STD006',
            'code': 'UL 1703',
            'title': 'Flat-Plate Photovoltaic Modules and Panels',
            'category': 'Safety',
            'authority': 'UL',
            'year': 2021,
            'citation': 'UL 1703:2021, Flat-Plate Photovoltaic Modules and Panels'
        },
        {
            'id': 'STD007',
            'code': 'ASTM E1171',
            'title': 'Standard Test Method for Photovoltaic Modules in Cyclic Temperature and Humidity Environments',
            'category': 'Environmental Testing',
            'authority': 'ASTM',
            'year': 2020,
            'citation': 'ASTM E1171-20, Standard Test Method for Photovoltaic Modules in Cyclic Temperature and Humidity Environments'
        },
        {
            'id': 'STD008',
            'code': 'ISO/IEC 17025:2017',
            'title': 'General requirements for the competence of testing and calibration laboratories',
            'category': 'Quality Management',
            'authority': 'ISO/IEC',
            'year': 2017,
            'citation': 'ISO/IEC 17025:2017, General requirements for the competence of testing and calibration laboratories'
        },
        {
            'id': 'STD009',
            'code': 'MNRE Guidelines 2023',
            'title': 'Ministry of New and Renewable Energy - Technical Standards for Grid Connected Rooftop Solar PV Systems',
            'category': 'Government Guidelines',
            'authority': 'MNRE',
            'year': 2023,
            'citation': 'MNRE Guidelines 2023, Technical Standards for Grid Connected Rooftop Solar PV Systems'
        }
    ]

def add_audit_trail(action: str, entity_type: str, entity_id: str, details: str = ""):
    """Add entry to audit trail"""
    st.session_state.audit_trail.append({
        'timestamp': datetime.now(),
        'user': st.session_state.current_user,
        'action': action,
        'entity_type': entity_type,
        'entity_id': entity_id,
        'details': details
    })

def create_notification(recipient: str, notification_type: str, title: str, message: str, priority: str = "Medium"):
    """Create a notification"""
    st.session_state.notifications.append({
        'id': str(uuid.uuid4()),
        'recipient': recipient,
        'type': notification_type,
        'title': title,
        'message': message,
        'priority': priority,
        'timestamp': datetime.now(),
        'read': False
    })

# ============================================================================
# TRANSLATION FUNCTIONS
# ============================================================================

SUPPORTED_LANGUAGES = {
    'English': 'en',
    'Hindi': 'hi',
    'Tamil': 'ta',
    'Telugu': 'te',
    'Gujarati': 'gu',
    'Marathi': 'mr',
    'Kannada': 'kn',
    'Spanish': 'es',
    'French': 'fr',
    'German': 'de',
    'Chinese': 'zh-CN',
    'Japanese': 'ja'
}

def translate_text(text: str, target_language: str) -> str:
    """Translate text to target language"""
    try:
        if target_language == 'en' or not text:
            return text
        translator = GoogleTranslator(source='en', target=target_language)
        return translator.translate(text)
    except Exception as e:
        st.warning(f"Translation error: {str(e)}")
        return text

def translate_document_section(section_content: Dict[str, Any], target_language: str) -> Dict[str, Any]:
    """Translate a document section"""
    translated = {}
    for key, value in section_content.items():
        if isinstance(value, str):
            translated[key] = translate_text(value, target_language)
        elif isinstance(value, list):
            translated[key] = [translate_text(item, target_language) if isinstance(item, str) else item for item in value]
        else:
            translated[key] = value
    return translated

# ============================================================================
# DOCUMENT GENERATION FUNCTIONS
# ============================================================================

def generate_word_document(sop_data: Dict[str, Any], translate_to: str = 'en') -> io.BytesIO:
    """Generate professional Word document with all SOP sections"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add custom styles
    styles = doc.styles

    # Title style
    title_style = styles['Heading 1']
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(255, 107, 53)

    # Heading 2 style
    heading2_style = styles['Heading 2']
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(0, 78, 137)

    # Document Header - Company Logo
    if sop_data.get('company_logo'):
        try:
            logo_stream = io.BytesIO(sop_data['company_logo'])
            doc.add_picture(logo_stream, width=Inches(1.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

    # Title
    title = doc.add_heading(translate_text(sop_data.get('title', 'Standard Operating Procedure'), translate_to), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document Information Table
    doc.add_paragraph()
    doc_info_table = doc.add_table(rows=6, cols=2)
    doc_info_table.style = 'Light Grid Accent 1'

    doc_info_data = [
        (translate_text('Document No.', translate_to), sop_data.get('doc_no', '')),
        (translate_text('Document Owner', translate_to), sop_data.get('doc_owner', '')),
        (translate_text('Division', translate_to), sop_data.get('division', '')),
        (translate_text('Company Name', translate_to), sop_data.get('company_name', '')),
        (translate_text('Effective Date', translate_to), sop_data.get('effective_date', '')),
        (translate_text('Revision No.', translate_to), sop_data.get('revision_no', '00')),
    ]

    for i, (label, value) in enumerate(doc_info_data):
        doc_info_table.rows[i].cells[0].text = label
        doc_info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        doc_info_table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # Doer-Reviewer-Approver Chain
    doc.add_heading(translate_text('Approval Chain', translate_to), level=2)
    approval_table = doc.add_table(rows=len(sop_data.get('approval_chain', [])) + 1, cols=4)
    approval_table.style = 'Light Grid Accent 1'

    # Header row
    headers = [translate_text(h, translate_to) for h in ['Role', 'Name', 'Signature', 'Date']]
    for i, header in enumerate(headers):
        approval_table.rows[0].cells[i].text = header
        approval_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    # Data rows
    for i, approver in enumerate(sop_data.get('approval_chain', []), 1):
        approval_table.rows[i].cells[0].text = translate_text(approver.get('role', ''), translate_to)
        approval_table.rows[i].cells[1].text = approver.get('name', '')
        approval_table.rows[i].cells[2].text = approver.get('signature', '[Signature]')
        approval_table.rows[i].cells[3].text = approver.get('date', '')

    doc.add_page_break()

    # Revision History
    doc.add_heading(translate_text('Revision History', translate_to), level=2)
    revision_table = doc.add_table(rows=len(sop_data.get('revision_history', [])) + 1, cols=4)
    revision_table.style = 'Light Grid Accent 1'

    # Header row
    rev_headers = [translate_text(h, translate_to) for h in ['Revision No.', 'Date', 'Description', 'Author']]
    for i, header in enumerate(rev_headers):
        revision_table.rows[0].cells[i].text = header
        revision_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    # Data rows
    for i, revision in enumerate(sop_data.get('revision_history', []), 1):
        revision_table.rows[i].cells[0].text = revision.get('rev_no', '')
        revision_table.rows[i].cells[1].text = revision.get('date', '')
        revision_table.rows[i].cells[2].text = translate_text(revision.get('description', ''), translate_to)
        revision_table.rows[i].cells[3].text = revision.get('author', '')

    doc.add_page_break()

    # Table of Contents
    doc.add_heading(translate_text('Table of Contents', translate_to), level=1)
    toc_items = [
        '1. Purpose',
        '2. Scope',
        '3. Definitions and Abbreviations',
        '4. Responsibilities',
        '5. Normative References',
        '6. HSE Risk Assessment',
        '7. Equipment and Materials Required',
        '8. Test Procedure',
        '9. Analysis Methodology',
        '10. Final Requirements',
        '11. Pass/Fail Criteria',
        '12. Test Schematic',
        '13. Appendix'
    ]
    for item in toc_items:
        p = doc.add_paragraph(translate_text(item, translate_to), style='List Number')

    doc.add_page_break()

    # 1. Purpose
    doc.add_heading(translate_text('1. Purpose', translate_to), level=1)
    doc.add_paragraph(translate_text(sop_data.get('purpose', 'This SOP defines the procedure for conducting the test.'), translate_to))
    doc.add_paragraph()

    # 2. Scope
    doc.add_heading(translate_text('2. Scope', translate_to), level=1)
    doc.add_paragraph(translate_text(sop_data.get('scope', 'This procedure applies to all relevant test activities.'), translate_to))
    doc.add_paragraph()

    # 3. Definitions
    doc.add_heading(translate_text('3. Definitions and Abbreviations', translate_to), level=1)
    for definition in sop_data.get('definitions', []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(translate_text(definition.get('term', ''), translate_to)).bold = True
        p.add_run(f": {translate_text(definition.get('definition', ''), translate_to)}")
    doc.add_paragraph()

    # 4. Responsibilities
    doc.add_heading(translate_text('4. Responsibilities', translate_to), level=1)
    resp_table = doc.add_table(rows=len(sop_data.get('responsibilities', [])) + 1, cols=2)
    resp_table.style = 'Light Grid Accent 1'
    resp_table.rows[0].cells[0].text = translate_text('Role', translate_to)
    resp_table.rows[0].cells[1].text = translate_text('Responsibility', translate_to)
    for i, header_cell in enumerate(resp_table.rows[0].cells):
        header_cell.paragraphs[0].runs[0].bold = True

    for i, resp in enumerate(sop_data.get('responsibilities', []), 1):
        resp_table.rows[i].cells[0].text = translate_text(resp.get('role', ''), translate_to)
        resp_table.rows[i].cells[1].text = translate_text(resp.get('responsibility', ''), translate_to)
    doc.add_paragraph()

    # 5. Normative References
    doc.add_heading(translate_text('5. Normative References', translate_to), level=1)
    for ref in sop_data.get('references', []):
        doc.add_paragraph(translate_text(ref, translate_to), style='List Bullet')
    doc.add_paragraph()

    # 6. HSE Risk Assessment
    doc.add_heading(translate_text('6. Health, Safety & Environment (HSE) Risk Assessment', translate_to), level=1)
    hse_table = doc.add_table(rows=len(sop_data.get('hse_risks', [])) + 1, cols=4)
    hse_table.style = 'Light Grid Accent 1'

    hse_headers = [translate_text(h, translate_to) for h in ['Hazard', 'Risk Level', 'Control Measures', 'PPE Required']]
    for i, header in enumerate(hse_headers):
        hse_table.rows[0].cells[i].text = header
        hse_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for i, risk in enumerate(sop_data.get('hse_risks', []), 1):
        hse_table.rows[i].cells[0].text = translate_text(risk.get('hazard', ''), translate_to)
        hse_table.rows[i].cells[1].text = translate_text(risk.get('risk_level', ''), translate_to)
        hse_table.rows[i].cells[2].text = translate_text(risk.get('control_measures', ''), translate_to)
        hse_table.rows[i].cells[3].text = translate_text(risk.get('ppe', ''), translate_to)
    doc.add_paragraph()

    # 7. Equipment and Materials
    doc.add_heading(translate_text('7. Equipment and Materials Required', translate_to), level=1)
    equip_table = doc.add_table(rows=len(sop_data.get('equipment_list', [])) + 1, cols=4)
    equip_table.style = 'Light Grid Accent 1'

    equip_headers = [translate_text(h, translate_to) for h in ['S.No.', 'Equipment/Material', 'Standard Specification', 'Actual Details']]
    for i, header in enumerate(equip_headers):
        equip_table.rows[0].cells[i].text = header
        equip_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for i, equip in enumerate(sop_data.get('equipment_list', []), 1):
        equip_table.rows[i].cells[0].text = str(i)
        equip_table.rows[i].cells[1].text = translate_text(equip.get('name', ''), translate_to)
        equip_table.rows[i].cells[2].text = translate_text(equip.get('standard_spec', ''), translate_to)
        equip_table.rows[i].cells[3].text = equip.get('actual_details', '')
    doc.add_paragraph()

    # 8. Test Procedure
    doc.add_heading(translate_text('8. Detailed Test Procedure', translate_to), level=1)
    for step in sop_data.get('procedure_steps', []):
        p = doc.add_paragraph(style='List Number')
        p.add_run(translate_text(step, translate_to))
    doc.add_paragraph()

    # 9. Analysis Methodology
    doc.add_heading(translate_text('9. Analysis Methodology', translate_to), level=1)
    doc.add_paragraph(translate_text(sop_data.get('analysis_method', 'Analysis shall be performed as per standard requirements.'), translate_to))
    doc.add_paragraph()

    # 10. Final Requirements
    doc.add_heading(translate_text('10. Final Requirements', translate_to), level=1)
    for req in sop_data.get('final_requirements', []):
        doc.add_paragraph(translate_text(req, translate_to), style='List Bullet')
    doc.add_paragraph()

    # 11. Pass/Fail Criteria
    doc.add_heading(translate_text('11. Pass/Fail Criteria', translate_to), level=1)
    criteria_table = doc.add_table(rows=len(sop_data.get('pass_fail_criteria', [])) + 1, cols=3)
    criteria_table.style = 'Light Grid Accent 1'

    criteria_headers = [translate_text(h, translate_to) for h in ['Parameter', 'Pass Criteria', 'Fail Criteria']]
    for i, header in enumerate(criteria_headers):
        criteria_table.rows[0].cells[i].text = header
        criteria_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for i, criteria in enumerate(sop_data.get('pass_fail_criteria', []), 1):
        criteria_table.rows[i].cells[0].text = translate_text(criteria.get('parameter', ''), translate_to)
        criteria_table.rows[i].cells[1].text = translate_text(criteria.get('pass_criteria', ''), translate_to)
        criteria_table.rows[i].cells[2].text = translate_text(criteria.get('fail_criteria', ''), translate_to)

    # Add test schematic if available
    if sop_data.get('test_schematic'):
        doc.add_page_break()
        doc.add_heading(translate_text('12. Test Schematic', translate_to), level=1)
        try:
            schematic_stream = io.BytesIO(sop_data['test_schematic'])
            doc.add_picture(schematic_stream, width=Inches(6))
        except:
            doc.add_paragraph(translate_text('[Test schematic diagram]', translate_to))

    # Appendix
    doc.add_page_break()
    doc.add_heading(translate_text('13. Appendix', translate_to), level=1)
    doc.add_paragraph(translate_text(sop_data.get('appendix', 'Additional information and supporting documents.'), translate_to))

    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"{sop_data.get('doc_no', '')} - {translate_text('Page', translate_to)} "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return doc_io

def generate_pdf_document(sop_data: Dict[str, Any], translate_to: str = 'en') -> io.BytesIO:
    """Generate PDF document"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#FF6B35'),
        spaceAfter=30,
        alignment=TA_CENTER
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#004E89'),
        spaceAfter=12,
        spaceBefore=12
    )

    # Title
    story.append(Paragraph(translate_text(sop_data.get('title', 'Standard Operating Procedure'), translate_to), title_style))
    story.append(Spacer(1, 12))

    # Document Info Table
    doc_info_data = [
        [translate_text('Document No.', translate_to), sop_data.get('doc_no', '')],
        [translate_text('Document Owner', translate_to), sop_data.get('doc_owner', '')],
        [translate_text('Division', translate_to), sop_data.get('division', '')],
        [translate_text('Company Name', translate_to), sop_data.get('company_name', '')],
    ]

    doc_info_table = Table(doc_info_data, colWidths=[3*inch, 3*inch])
    doc_info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E8F4F8')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))

    story.append(doc_info_table)
    story.append(PageBreak())

    # Purpose
    story.append(Paragraph(translate_text('1. Purpose', translate_to), heading_style))
    story.append(Paragraph(translate_text(sop_data.get('purpose', ''), translate_to), styles['Normal']))
    story.append(Spacer(1, 12))

    # Scope
    story.append(Paragraph(translate_text('2. Scope', translate_to), heading_style))
    story.append(Paragraph(translate_text(sop_data.get('scope', ''), translate_to), styles['Normal']))
    story.append(Spacer(1, 12))

    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_excel_document(sop_data: Dict[str, Any]) -> io.BytesIO:
    """Generate Excel document with SOP data"""
    buffer = io.BytesIO()
    workbook = openpyxl.Workbook()

    # Document Info Sheet
    ws_info = workbook.active
    ws_info.title = "Document Info"

    # Header styling
    header_fill = PatternFill(start_color="FF6B35", end_color="FF6B35", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)

    ws_info['A1'] = 'Document Information'
    ws_info['A1'].font = Font(bold=True, size=14)
    ws_info.merge_cells('A1:B1')

    info_data = [
        ['Document No.', sop_data.get('doc_no', '')],
        ['Title', sop_data.get('title', '')],
        ['Document Owner', sop_data.get('doc_owner', '')],
        ['Division', sop_data.get('division', '')],
        ['Company Name', sop_data.get('company_name', '')],
        ['Effective Date', sop_data.get('effective_date', '')],
        ['Revision No.', sop_data.get('revision_no', '00')],
    ]

    for i, (label, value) in enumerate(info_data, 3):
        ws_info.cell(row=i, column=1, value=label).font = Font(bold=True)
        ws_info.cell(row=i, column=2, value=value)

    # Equipment Sheet
    ws_equip = workbook.create_sheet("Equipment List")
    ws_equip.append(['S.No.', 'Equipment/Material', 'Standard Specification', 'Actual Details'])

    for i, equip in enumerate(sop_data.get('equipment_list', []), 1):
        ws_equip.append([
            i,
            equip.get('name', ''),
            equip.get('standard_spec', ''),
            equip.get('actual_details', '')
        ])

    # Procedure Sheet
    ws_proc = workbook.create_sheet("Test Procedure")
    ws_proc.append(['Step No.', 'Procedure Step'])

    for i, step in enumerate(sop_data.get('procedure_steps', []), 1):
        ws_proc.append([i, step])

    # Pass/Fail Criteria Sheet
    ws_criteria = workbook.create_sheet("Pass-Fail Criteria")
    ws_criteria.append(['Parameter', 'Pass Criteria', 'Fail Criteria'])

    for criteria in sop_data.get('pass_fail_criteria', []):
        ws_criteria.append([
            criteria.get('parameter', ''),
            criteria.get('pass_criteria', ''),
            criteria.get('fail_criteria', '')
        ])

    workbook.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================================================
# UI RENDERING FUNCTIONS - SOP GENERATOR
# ============================================================================

def render_sop_generator():
    """Main SOP Generator Interface"""
    st.markdown('<h1 class="main-header">☀️ Solar PV SOP Generator</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Complete Standard Operating Procedure Document Generator</p>', unsafe_allow_html=True)

    # Tabs for different sections
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📝 Create SOP",
        "🌐 Translation",
        "📚 Standards Library",
        "📄 My Documents",
        "⚙️ Settings"
    ])

    with tab1:
        render_sop_creation()

    with tab2:
        render_translation_interface()

    with tab3:
        render_standards_library()

    with tab4:
        render_document_management()

    with tab5:
        render_settings()

def render_sop_creation():
    """SOP Creation Interface"""
    st.subheader("📝 Create New SOP Document")

    # Template Selection
    col1, col2 = st.columns([2, 1])
    with col1:
        template_options = ['Custom'] + [t['name'] for t in st.session_state.sop_templates]
        selected_template = st.selectbox("Select Template", template_options)
    with col2:
        st.metric("Total SOPs Created", len(st.session_state.sop_documents))

    st.markdown("---")

    # Create form for SOP
    with st.form("sop_creation_form"):
        st.markdown("### 📋 Document Header Information")

        col1, col2, col3 = st.columns(3)
        with col1:
            doc_title = st.text_input("Document Title *", placeholder="e.g., Thermal Cycling Test SOP")
            doc_no = st.text_input("Document Number *", placeholder="e.g., SOP-TC-001")
        with col2:
            doc_owner = st.text_input("Document Owner *", placeholder="e.g., Dr. Rajesh Kumar")
            division = st.text_input("Division", placeholder="e.g., Testing & Certification")
        with col3:
            company_name = st.text_input("Company Name", placeholder="e.g., Solar Test Lab Pvt Ltd")
            revision_no = st.text_input("Revision Number", value="00")

        effective_date = st.date_input("Effective Date", value=date.today())

        # Logo Upload
        st.markdown("### 🖼️ Company Logo Upload")
        company_logo = st.file_uploader("Upload Company Logo (PNG/JPG)", type=['png', 'jpg', 'jpeg'], key="logo_upload")

        st.markdown("---")
        st.markdown("### 👥 Approval Chain (Doer-Reviewer-Approver)")

        # Dynamic approval chain
        num_approvers = st.number_input("Number of Approvers", min_value=1, max_value=10, value=3)

        approval_chain = []
        for i in range(num_approvers):
            col1, col2, col3 = st.columns(3)
            with col1:
                role = st.text_input(f"Role {i+1}", value=["Doer", "Reviewer", "Approver"][i] if i < 3 else f"Approver {i}", key=f"role_{i}")
            with col2:
                name = st.text_input(f"Name {i+1}", key=f"name_{i}")
            with col3:
                signature = st.text_input(f"Signature {i+1}", placeholder="[Signature]", key=f"sig_{i}")

            approval_chain.append({
                'role': role,
                'name': name,
                'signature': signature,
                'date': str(date.today())
            })

        st.markdown("---")
        st.markdown("### 📜 Revision History")

        num_revisions = st.number_input("Number of Revisions", min_value=1, max_value=10, value=1)

        revision_history = []
        for i in range(num_revisions):
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                rev_no = st.text_input(f"Rev {i+1}", value=f"{i:02d}", key=f"rev_no_{i}")
            with col2:
                rev_date = st.date_input(f"Date {i+1}", value=date.today(), key=f"rev_date_{i}")
            with col3:
                rev_desc = st.text_input(f"Description {i+1}", value="Initial Release" if i == 0 else "", key=f"rev_desc_{i}")
            with col4:
                rev_author = st.text_input(f"Author {i+1}", key=f"rev_author_{i}")

            revision_history.append({
                'rev_no': rev_no,
                'date': str(rev_date),
                'description': rev_desc,
                'author': rev_author
            })

        st.markdown("---")
        st.markdown("### 📖 Document Sections")

        # Purpose
        purpose = st.text_area("1. Purpose *",
                               placeholder="State the purpose of this test procedure...",
                               height=100)

        # Scope
        scope = st.text_area("2. Scope *",
                            placeholder="Define the scope and applicability...",
                            height=100)

        # Definitions
        st.markdown("**3. Definitions and Abbreviations**")
        num_definitions = st.number_input("Number of Definitions", min_value=0, max_value=20, value=3)

        definitions = []
        for i in range(num_definitions):
            col1, col2 = st.columns(2)
            with col1:
                term = st.text_input(f"Term {i+1}", key=f"def_term_{i}")
            with col2:
                definition = st.text_input(f"Definition {i+1}", key=f"def_def_{i}")
            definitions.append({'term': term, 'definition': definition})

        # Responsibilities
        st.markdown("**4. Responsibilities**")
        num_responsibilities = st.number_input("Number of Roles", min_value=1, max_value=10, value=3)

        responsibilities = []
        for i in range(num_responsibilities):
            col1, col2 = st.columns(2)
            with col1:
                resp_role = st.text_input(f"Role {i+1}", key=f"resp_role_{i}")
            with col2:
                resp_detail = st.text_input(f"Responsibility {i+1}", key=f"resp_detail_{i}")
            responsibilities.append({'role': resp_role, 'responsibility': resp_detail})

        # Normative References
        st.markdown("**5. Normative References**")

        # Standards selection from library
        available_standards = [std['citation'] for std in st.session_state.standards_library]
        selected_standards = st.multiselect("Select Standards from Library", available_standards)

        additional_refs = st.text_area("Additional References (one per line)",
                                       placeholder="Enter additional references...",
                                       height=100)

        references = list(selected_standards)
        if additional_refs:
            references.extend([ref.strip() for ref in additional_refs.split('\n') if ref.strip()])

        # HSE Risk Assessment
        st.markdown("**6. HSE Risk Assessment**")
        num_hse_risks = st.number_input("Number of HSE Risks", min_value=1, max_value=15, value=3)

        hse_risks = []
        for i in range(num_hse_risks):
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                hazard = st.text_input(f"Hazard {i+1}", key=f"hse_hazard_{i}")
            with col2:
                risk_level = st.selectbox(f"Risk Level {i+1}", ['Low', 'Medium', 'High', 'Critical'], key=f"hse_level_{i}")
            with col3:
                control = st.text_input(f"Control Measures {i+1}", key=f"hse_control_{i}")
            with col4:
                ppe = st.text_input(f"PPE Required {i+1}", key=f"hse_ppe_{i}")

            hse_risks.append({
                'hazard': hazard,
                'risk_level': risk_level,
                'control_measures': control,
                'ppe': ppe
            })

        # Equipment and Materials
        st.markdown("**7. Equipment and Materials Required**")
        num_equipment = st.number_input("Number of Equipment/Materials", min_value=1, max_value=20, value=3)

        equipment_list = []
        for i in range(num_equipment):
            col1, col2, col3 = st.columns(3)
            with col1:
                equip_name = st.text_input(f"Equipment {i+1}", key=f"equip_name_{i}")
            with col2:
                equip_standard = st.text_input(f"Standard Spec {i+1}", key=f"equip_std_{i}")
            with col3:
                equip_actual = st.text_input(f"Actual Details {i+1}", key=f"equip_actual_{i}")

            equipment_list.append({
                'name': equip_name,
                'standard_spec': equip_standard,
                'actual_details': equip_actual
            })

        # Equipment Photos Upload
        equipment_photos = st.file_uploader("Upload Equipment Photos (Multiple allowed)",
                                           type=['png', 'jpg', 'jpeg'],
                                           accept_multiple_files=True,
                                           key="equipment_photos")

        # Test Procedure Steps
        st.markdown("**8. Detailed Test Procedure**")
        num_steps = st.number_input("Number of Procedure Steps", min_value=1, max_value=50, value=10)

        procedure_steps = []
        for i in range(num_steps):
            step = st.text_input(f"Step {i+1}", key=f"proc_step_{i}")
            if step:
                procedure_steps.append(step)

        # Flowchart Upload
        flowchart = st.file_uploader("Upload Process Flowchart (Optional)",
                                    type=['png', 'jpg', 'jpeg', 'pdf'],
                                    key="flowchart")

        # Analysis Methodology
        analysis_method = st.text_area("9. Analysis Methodology",
                                       placeholder="Describe how the test results will be analyzed...",
                                       height=150)

        # Final Requirements
        st.markdown("**10. Final Requirements**")
        final_reqs_text = st.text_area("Final Requirements (one per line)",
                                       placeholder="Enter final requirements...",
                                       height=100)

        final_requirements = [req.strip() for req in final_reqs_text.split('\n') if req.strip()]

        # Pass/Fail Criteria
        st.markdown("**11. Pass/Fail Criteria**")
        num_criteria = st.number_input("Number of Criteria", min_value=1, max_value=20, value=3)

        pass_fail_criteria = []
        for i in range(num_criteria):
            col1, col2, col3 = st.columns(3)
            with col1:
                param = st.text_input(f"Parameter {i+1}", key=f"criteria_param_{i}")
            with col2:
                pass_crit = st.text_input(f"Pass Criteria {i+1}", key=f"criteria_pass_{i}")
            with col3:
                fail_crit = st.text_input(f"Fail Criteria {i+1}", key=f"criteria_fail_{i}")

            pass_fail_criteria.append({
                'parameter': param,
                'pass_criteria': pass_crit,
                'fail_criteria': fail_crit
            })

        # Test Schematic
        test_schematic = st.file_uploader("Upload Test Schematic Diagram",
                                         type=['png', 'jpg', 'jpeg', 'pdf'],
                                         key="test_schematic")

        # Data Tables Upload
        data_tables = st.file_uploader("Upload Data Tables (CSV/Excel)",
                                       type=['csv', 'xlsx'],
                                       accept_multiple_files=True,
                                       key="data_tables")

        # Appendix
        appendix = st.text_area("13. Appendix",
                               placeholder="Additional information, supporting documents, references...",
                               height=150)

        st.markdown("---")
        st.markdown("### 🌐 Translation Options")

        translate_to_lang = st.selectbox("Translate Document To", ['English'] + list(SUPPORTED_LANGUAGES.keys())[1:])

        st.markdown("---")
        st.markdown("### 📥 Export Options")

        col1, col2, col3 = st.columns(3)
        with col1:
            export_word = st.checkbox("Export as Word (.docx)", value=True)
        with col2:
            export_pdf = st.checkbox("Export as PDF")
        with col3:
            export_excel = st.checkbox("Export as Excel")

        # Submit button
        submitted = st.form_submit_button("🚀 Generate SOP Document", use_container_width=True)

        if submitted:
            # Validation
            if not doc_title or not doc_no or not doc_owner or not purpose or not scope:
                st.error("Please fill all required fields marked with *")
            else:
                # Prepare SOP data
                sop_data = {
                    'id': str(uuid.uuid4()),
                    'title': doc_title,
                    'doc_no': doc_no,
                    'doc_owner': doc_owner,
                    'division': division,
                    'company_name': company_name,
                    'effective_date': str(effective_date),
                    'revision_no': revision_no,
                    'approval_chain': approval_chain,
                    'revision_history': revision_history,
                    'purpose': purpose,
                    'scope': scope,
                    'definitions': definitions,
                    'responsibilities': responsibilities,
                    'references': references,
                    'hse_risks': hse_risks,
                    'equipment_list': equipment_list,
                    'procedure_steps': procedure_steps,
                    'analysis_method': analysis_method,
                    'final_requirements': final_requirements,
                    'pass_fail_criteria': pass_fail_criteria,
                    'appendix': appendix,
                    'created_date': datetime.now(),
                    'created_by': st.session_state.current_user,
                    'template_used': selected_template
                }

                # Store uploaded files
                if company_logo:
                    sop_data['company_logo'] = company_logo.read()
                    company_logo.seek(0)

                if test_schematic:
                    sop_data['test_schematic'] = test_schematic.read()
                    test_schematic.seek(0)

                # Save to session state
                st.session_state.sop_documents.append(sop_data)

                # Generate documents
                with st.spinner("Generating SOP documents..."):
                    progress_bar = st.progress(0)

                    # Get translation language code
                    translate_lang_code = SUPPORTED_LANGUAGES.get(translate_to_lang, 'en')

                    # Word Document
                    if export_word:
                        progress_bar.progress(33)
                        word_doc = generate_word_document(sop_data, translate_lang_code)
                        st.download_button(
                            label="📄 Download Word Document",
                            data=word_doc,
                            file_name=f"{doc_no}_SOP.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                    # PDF Document
                    if export_pdf:
                        progress_bar.progress(66)
                        pdf_doc = generate_pdf_document(sop_data, translate_lang_code)
                        st.download_button(
                            label="📕 Download PDF Document",
                            data=pdf_doc,
                            file_name=f"{doc_no}_SOP.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )

                    # Excel Document
                    if export_excel:
                        progress_bar.progress(100)
                        excel_doc = generate_excel_document(sop_data)
                        st.download_button(
                            label="📊 Download Excel Document",
                            data=excel_doc,
                            file_name=f"{doc_no}_SOP.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )

                    progress_bar.progress(100)

                # Add to audit trail
                add_audit_trail("SOP Created", "SOP Document", sop_data['id'], f"Created SOP: {doc_title}")

                st.success(f"✅ SOP Document '{doc_title}' created successfully!")
                st.balloons()

def render_translation_interface():
    """Translation Interface for existing documents"""
    st.subheader("🌐 Document Translation")

    if not st.session_state.sop_documents:
        st.info("📝 No SOP documents available. Create one first in the 'Create SOP' tab.")
        return

    # Select document to translate
    doc_options = [f"{doc['doc_no']} - {doc['title']}" for doc in st.session_state.sop_documents]
    selected_doc_idx = st.selectbox("Select Document to Translate", range(len(doc_options)), format_func=lambda x: doc_options[x])

    if selected_doc_idx is not None:
        selected_doc = st.session_state.sop_documents[selected_doc_idx]

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Original Language:** English")
        with col2:
            target_language = st.selectbox("Translate To", list(SUPPORTED_LANGUAGES.keys())[1:])

        # Section selection
        sections_to_translate = st.multiselect(
            "Select Sections to Translate",
            ['All Sections', 'Purpose', 'Scope', 'Definitions', 'Responsibilities',
             'HSE Risks', 'Equipment List', 'Procedure Steps', 'Analysis',
             'Final Requirements', 'Pass/Fail Criteria', 'Appendix'],
            default=['All Sections']
        )

        if st.button("🌐 Translate Document", use_container_width=True):
            with st.spinner(f"Translating to {target_language}..."):
                target_lang_code = SUPPORTED_LANGUAGES[target_language]

                # Generate translated document
                word_doc = generate_word_document(selected_doc, target_lang_code)

                st.success(f"✅ Translation to {target_language} completed!")

                st.download_button(
                    label=f"📄 Download Translated Document ({target_language})",
                    data=word_doc,
                    file_name=f"{selected_doc['doc_no']}_SOP_{target_language}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

def render_standards_library():
    """Standards Library Management"""
    st.subheader("📚 Standards & Citations Library")

    # Display standards
    df_standards = pd.DataFrame(st.session_state.standards_library)

    # Filter by category
    categories = ['All'] + list(df_standards['category'].unique())
    selected_category = st.selectbox("Filter by Category", categories)

    if selected_category != 'All':
        df_filtered = df_standards[df_standards['category'] == selected_category]
    else:
        df_filtered = df_standards

    # Display as cards
    for idx, std in df_filtered.iterrows():
        with st.expander(f"📖 {std['code']} - {std['title'][:80]}..."):
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**Full Title:** {std['title']}")
                st.markdown(f"**Category:** {std['category']}")
                st.markdown(f"**Authority:** {std['authority']}")
                st.markdown(f"**Year:** {std['year']}")
                st.markdown(f"**Citation:**")
                st.code(std['citation'])
            with col2:
                if st.button("📋 Copy Citation", key=f"copy_{std['id']}"):
                    st.success("Citation copied!")

    st.markdown("---")

    # Add new standard
    with st.expander("➕ Add New Standard"):
        with st.form("add_standard_form"):
            col1, col2 = st.columns(2)
            with col1:
                new_code = st.text_input("Standard Code", placeholder="e.g., IEC 61215-1:2016")
                new_authority = st.selectbox("Authority", ['IEC', 'ISO', 'IS', 'UL', 'ASTM', 'EN', 'MNRE', 'NABL', 'Other'])
            with col2:
                new_category = st.selectbox("Category",
                    ['PV Module Testing', 'Safety', 'Environmental Testing', 'Quality Management',
                     'Indian Standard', 'Government Guidelines', 'Other'])
                new_year = st.number_input("Year", min_value=1950, max_value=2030, value=2023)

            new_title = st.text_area("Full Title", placeholder="Enter the complete title of the standard...")
            new_citation = st.text_area("Citation Format", placeholder="Enter how this standard should be cited...")

            if st.form_submit_button("Add Standard"):
                if new_code and new_title:
                    new_standard = {
                        'id': f"STD{len(st.session_state.standards_library) + 1:03d}",
                        'code': new_code,
                        'title': new_title,
                        'category': new_category,
                        'authority': new_authority,
                        'year': new_year,
                        'citation': new_citation if new_citation else f"{new_code}, {new_title}"
                    }
                    st.session_state.standards_library.append(new_standard)
                    st.success(f"✅ Standard {new_code} added successfully!")
                    st.rerun()

def render_document_management():
    """Document Management Interface"""
    st.subheader("📄 My SOP Documents")

    if not st.session_state.sop_documents:
        st.info("📝 No SOP documents created yet. Create your first SOP in the 'Create SOP' tab.")
        return

    # Display documents
    for idx, doc in enumerate(st.session_state.sop_documents):
        with st.expander(f"📄 {doc['doc_no']} - {doc['title']}"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"**Document Owner:** {doc['doc_owner']}")
                st.markdown(f"**Division:** {doc.get('division', 'N/A')}")
            with col2:
                st.markdown(f"**Revision:** {doc.get('revision_no', '00')}")
                st.markdown(f"**Effective Date:** {doc.get('effective_date', 'N/A')}")
            with col3:
                st.markdown(f"**Created By:** {doc.get('created_by', 'N/A')}")
                st.markdown(f"**Created On:** {doc['created_date'].strftime('%Y-%m-%d %H:%M')}")

            st.markdown("---")

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                if st.button("📄 Export Word", key=f"word_{idx}"):
                    word_doc = generate_word_document(doc)
                    st.download_button(
                        label="Download",
                        data=word_doc,
                        file_name=f"{doc['doc_no']}_SOP.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_word_{idx}"
                    )
            with col2:
                if st.button("📕 Export PDF", key=f"pdf_{idx}"):
                    pdf_doc = generate_pdf_document(doc)
                    st.download_button(
                        label="Download",
                        data=pdf_doc,
                        file_name=f"{doc['doc_no']}_SOP.pdf",
                        mime="application/pdf",
                        key=f"dl_pdf_{idx}"
                    )
            with col3:
                if st.button("📊 Export Excel", key=f"excel_{idx}"):
                    excel_doc = generate_excel_document(doc)
                    st.download_button(
                        label="Download",
                        data=excel_doc,
                        file_name=f"{doc['doc_no']}_SOP.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key=f"dl_excel_{idx}"
                    )
            with col4:
                if st.button("🗑️ Delete", key=f"delete_{idx}"):
                    st.session_state.sop_documents.pop(idx)
                    st.success("Document deleted!")
                    st.rerun()

def render_settings():
    """Settings and Configuration"""
    st.subheader("⚙️ Settings & Configuration")

    st.markdown("### 👤 User Information")
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.current_user = st.text_input("Current User", value=st.session_state.current_user)
    with col2:
        st.session_state.user_role = st.selectbox("User Role",
            ['Engineer', 'Senior Engineer', 'Manager', 'Quality Manager', 'Lab Director'],
            index=['Engineer', 'Senior Engineer', 'Manager', 'Quality Manager', 'Lab Director'].index(st.session_state.user_role))

    st.markdown("---")
    st.markdown("### 📊 Statistics")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total SOPs", len(st.session_state.sop_documents))
    with col2:
        st.metric("Standards Library", len(st.session_state.standards_library))
    with col3:
        st.metric("Templates", len(st.session_state.sop_templates))
    with col4:
        st.metric("Audit Entries", len(st.session_state.audit_trail))

    st.markdown("---")
    st.markdown("### 🔄 Data Management")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("📥 Export All Data (JSON)", use_container_width=True):
            export_data = {
                'sop_documents': st.session_state.sop_documents,
                'standards_library': st.session_state.standards_library,
                'sop_templates': st.session_state.sop_templates,
                'audit_trail': [
                    {**entry, 'timestamp': entry['timestamp'].isoformat() if isinstance(entry['timestamp'], datetime) else entry['timestamp']}
                    for entry in st.session_state.audit_trail
                ]
            }

            # Convert to JSON
            json_str = json.dumps(export_data, indent=2, default=str)

            st.download_button(
                label="Download JSON",
                data=json_str,
                file_name=f"sop_data_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )

    with col2:
        if st.button("🗑️ Clear All Data", use_container_width=True):
            if st.checkbox("I understand this will delete all data"):
                st.session_state.sop_documents = []
                st.session_state.audit_trail = []
                st.success("All data cleared!")
                st.rerun()

# ============================================================================
# INTEGRATED MODULES FROM EXISTING APP
# ============================================================================

def render_dashboard():
    """Dashboard with metrics and charts"""
    st.markdown('<h2 class="sub-header">📊 Dashboard Overview</h2>', unsafe_allow_html=True)

    # Metrics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Active Projects", len([p for p in st.session_state.projects if p['status'] == 'Active']))
    with col2:
        st.metric("Total Samples", len(st.session_state.samples))
    with col3:
        st.metric("Equipment Available", len([e for e in st.session_state.equipment if e['status'] == 'Available']))
    with col4:
        st.metric("SOP Documents", len(st.session_state.sop_documents))

    st.markdown("---")

    # Charts
    col1, col2 = st.columns(2)

    with col1:
        # Project Status Distribution
        if st.session_state.projects:
            status_counts = pd.DataFrame(st.session_state.projects)['status'].value_counts()
            fig = px.pie(values=status_counts.values, names=status_counts.index,
                        title="Project Status Distribution")
            st.plotly_chart(fig, use_container_width=True)

    with col2:
        # Equipment Utilization
        if st.session_state.equipment:
            equip_df = pd.DataFrame(st.session_state.equipment)
            fig = px.bar(equip_df, x='name', y='utilization',
                        title="Equipment Utilization (%)")
            st.plotly_chart(fig, use_container_width=True)

def render_project_management():
    """Project Management Interface"""
    st.markdown('<h2 class="sub-header">📋 Project Management</h2>', unsafe_allow_html=True)

    # Display projects
    for project in st.session_state.projects:
        with st.expander(f"📁 {project['name']} ({project['id']})"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"**Client:** {project['client']}")
                st.markdown(f"**Manager:** {project['manager']}")
            with col2:
                st.markdown(f"**Status:** {project['status']}")
                st.markdown(f"**Priority:** {project['priority']}")
            with col3:
                st.markdown(f"**Budget:** ${project['budget']:,.2f}")
                st.markdown(f"**Spent:** ${project['spent']:,.2f}")

            # Progress bar
            progress = (project['spent'] / project['budget']) * 100 if project['budget'] > 0 else 0
            st.progress(min(progress / 100, 1.0))
            st.caption(f"{progress:.1f}% of budget used")

def render_equipment_management():
    """Equipment Management Interface"""
    st.markdown('<h2 class="sub-header">🔧 Equipment Management</h2>', unsafe_allow_html=True)

    for equipment in st.session_state.equipment:
        with st.expander(f"🔧 {equipment['name']} ({equipment['id']})"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"**Type:** {equipment['type']}")
                st.markdown(f"**Status:** {equipment['status']}")
                st.markdown(f"**Location:** {equipment['location']}")
            with col2:
                st.markdown(f"**Manufacturer:** {equipment.get('manufacturer', 'N/A')}")
                st.markdown(f"**Model:** {equipment.get('model', 'N/A')}")
                st.markdown(f"**Performance:** {equipment['performance_metric']}%")
            with col3:
                st.markdown(f"**Last Calibration:** {equipment['last_calibration']}")
                st.markdown(f"**Next Calibration:** {equipment['next_calibration']}")
                st.markdown(f"**Utilization:** {equipment['utilization']}%")

def render_manpower_management():
    """Manpower Management Interface"""
    st.markdown('<h2 class="sub-header">👥 Manpower Management</h2>', unsafe_allow_html=True)

    for person in st.session_state.manpower:
        with st.expander(f"👤 {person['name']} - {person['role']}"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"**Role:** {person['role']}")
                st.markdown(f"**Availability:** {person['availability']}")
                st.markdown(f"**Performance Score:** {person['performance_score']}%")
            with col2:
                st.markdown(f"**Skills:**")
                for skill in person['skills']:
                    st.markdown(f"- {skill}")
            with col3:
                st.markdown(f"**Certifications:**")
                for cert in person['certifications']:
                    st.markdown(f"- {cert}")
                st.markdown(f"**Hours Logged:** {person['hours_logged']}")

# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    """Main application"""
    # Initialize session state
    init_session_state()

    # Sidebar navigation
    st.sidebar.image("https://img.icons8.com/color/96/000000/solar-panel.png", width=100)
    st.sidebar.title("☀️ Solar PV Lab")
    st.sidebar.markdown("---")

    # Main menu
    menu_options = [
        "🏠 Home",
        "📝 SOP Generator",
        "📊 Dashboard",
        "📋 Projects",
        "🔧 Equipment",
        "👥 Manpower",
        "📄 Reports"
    ]

    selected_menu = st.sidebar.radio("Navigation", menu_options)

    st.sidebar.markdown("---")
    st.sidebar.markdown(f"**User:** {st.session_state.current_user}")
    st.sidebar.markdown(f"**Role:** {st.session_state.user_role}")

    # Render selected page
    if selected_menu == "🏠 Home":
        st.markdown('<h1 class="main-header">☀️ Solar PV Reliability Test Lab</h1>', unsafe_allow_html=True)
        st.markdown('<p class="sub-header">Complete Project Management & SOP Generation System</p>', unsafe_allow_html=True)

        st.markdown("""
        ### Welcome to the Solar PV Lab Management System

        This comprehensive platform integrates:
        - **SOP Generator**: Create professional Standard Operating Procedures
        - **Project Management**: Track projects, tasks, and WBS
        - **Equipment Management**: Monitor equipment calibration and utilization
        - **Manpower Management**: Manage team resources and schedules
        - **Translation**: Multi-language document support
        - **Standards Library**: Built-in citation management
        - **Document Export**: Word, PDF, and Excel formats

        Navigate using the sidebar menu to access different modules.
        """)

        st.markdown("---")
        render_dashboard()

    elif selected_menu == "📝 SOP Generator":
        render_sop_generator()

    elif selected_menu == "📊 Dashboard":
        render_dashboard()

    elif selected_menu == "📋 Projects":
        render_project_management()

    elif selected_menu == "🔧 Equipment":
        render_equipment_management()

    elif selected_menu == "👥 Manpower":
        render_manpower_management()

    elif selected_menu == "📄 Reports":
        st.markdown('<h2 class="sub-header">📄 Reports & Documents</h2>', unsafe_allow_html=True)
        st.info("Report generation module - integrated with SOP Generator")

if __name__ == "__main__":
    main()
